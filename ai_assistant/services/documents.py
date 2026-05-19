"""
Extracción de texto y contexto para documentos del chat IA.
"""
import csv
import io
import logging
import os
from pathlib import Path

logger = logging.getLogger(__name__)

ALLOWED_EXTENSIONS = {
    '.pdf', '.docx', '.doc', '.xlsx', '.xls',
    '.csv', '.txt', '.md', '.rtf',
}

EXTENSION_LABELS = {
    '.pdf': 'PDF',
    '.docx': 'Word',
    '.doc': 'Word',
    '.xlsx': 'Excel',
    '.xls': 'Excel',
    '.csv': 'CSV',
    '.txt': 'Texto',
    '.md': 'Markdown',
    '.rtf': 'RTF',
}

# Límite al guardar en BD (extracción)
MAX_EXTRACTED_CHARS = 8000
# Límite al enviar a la API (Groq rechaza payloads muy grandes → 413)
MAX_API_DOCUMENT_CHARS = 4500
MAX_PER_DOC_API_CHARS = 4000
MAX_PDF_PAGES = 20


def get_extension(filename):
    return Path(filename).suffix.lower()


def is_allowed_extension(filename):
    return get_extension(filename) in ALLOWED_EXTENSIONS


def extract_text_from_file(uploaded_file):
    """
    Extrae texto plano de un archivo subido.
    Returns: (text, error_message) — error_message es None si OK.
    """
    name = uploaded_file.name
    ext = get_extension(name)

    if not is_allowed_extension(name):
        return '', f'Tipo de archivo no permitido ({ext or "sin extensión"}).'

    try:
        uploaded_file.seek(0)
        raw = uploaded_file.read()
        uploaded_file.seek(0)
    except Exception as e:
        return '', f'No se pudo leer el archivo: {e}'

    try:
        if ext in ('.txt', '.md', '.csv', '.rtf'):
            text = _extract_plain(raw, ext)
        elif ext == '.pdf':
            text = _extract_pdf(raw)
        elif ext in ('.docx', '.doc'):
            text = _extract_word(raw, ext)
        elif ext in ('.xlsx', '.xls'):
            text = _extract_spreadsheet(raw, ext, name)
        else:
            return '', 'Formato no soportado.'
    except Exception as e:
        logger.exception('Error extrayendo texto de %s', name)
        return '', f'No se pudo procesar el archivo: {e}'

    text = _normalize_text(text)
    if not text.strip():
        return '', 'El archivo no contiene texto legible o está vacío.'

    if len(text) > MAX_EXTRACTED_CHARS:
        text = (
            text[:MAX_EXTRACTED_CHARS]
            + '\n\n[... resto del archivo omitido al extraer (documento muy largo) ...]'
        )

    return text, None


def truncate_text(text, max_chars, label='contenido'):
    """Recorta texto para caber en el límite de la API."""
    text = (text or '').strip()
    if len(text) <= max_chars:
        return text, False
    half = max_chars // 2
    return (
        text[:half]
        + f'\n\n[... {label} recortado por límite de la API ({len(text)} caracteres) ...]\n\n'
        + text[-half:]
    ), True


def _normalize_text(text):
    if not text:
        return ''
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    lines = [line.rstrip() for line in text.split('\n')]
    return '\n'.join(lines).strip()


def _extract_plain(raw, ext):
    for encoding in ('utf-8', 'utf-8-sig', 'latin-1', 'cp1252'):
        try:
            text = raw.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        text = raw.decode('utf-8', errors='replace')

    if ext == '.csv':
        return _csv_to_text(text)
    return text


def _csv_to_text(text):
    reader = csv.reader(io.StringIO(text))
    rows = []
    for i, row in enumerate(reader):
        if i > 500:
            rows.append('[... más filas omitidas ...]')
            break
        rows.append(' | '.join(cell.strip() for cell in row))
    return '\n'.join(rows)


def _extract_pdf(raw):
    try:
        from pypdf import PdfReader
    except ImportError:
        from PyPDF2 import PdfReader

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for i, page in enumerate(reader.pages):
        if i >= MAX_PDF_PAGES:
            parts.append(f'[... más de {MAX_PDF_PAGES} páginas omitidas ...]')
            break
        page_text = page.extract_text() or ''
        if page_text.strip():
            parts.append(page_text)
    return '\n\n'.join(parts)


def _extract_word(raw, ext):
    if ext == '.doc':
        return (
            '[Archivo .doc antiguo]\n'
            'Guarda el documento como .docx en Word y vuelve a subirlo.'
        )
    try:
        import docx
    except ImportError as e:
        raise RuntimeError('Falta la librería python-docx en el servidor.') from e

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(' | '.join(cells))
    return '\n'.join(parts)


def _extract_spreadsheet(raw, ext, filename):
    if ext == '.xls':
        return (
            '[Archivo .xls antiguo]\n'
            'Guarda la hoja como .xlsx en Excel y vuelve a subirla.'
        )
    try:
        import openpyxl
    except ImportError as e:
        raise RuntimeError('Falta la librería openpyxl en el servidor.') from e

    wb = openpyxl.load_workbook(io.BytesIO(raw), read_only=True, data_only=True)
    parts = []
    for sheet_name in wb.sheetnames[:5]:
        ws = wb[sheet_name]
        parts.append(f'--- Hoja: {sheet_name} ---')
        row_count = 0
        for row in ws.iter_rows(values_only=True):
            if row_count >= 500:
                parts.append('[... más filas omitidas ...]')
                break
            cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
            if cells:
                parts.append(' | '.join(cells))
                row_count += 1
    wb.close()
    return '\n'.join(parts)


def build_documents_context(documents, for_api=True):
    """
    Construye bloque de contexto para inyectar en el mensaje del usuario.
    documents: iterable de objetos con original_name y extracted_text.
  for_api: aplica límites estrictos para evitar error 413 en Groq.
    """
    if not documents:
        return '', False

    blocks = []
    truncated_any = False
    docs_list = list(documents)
    per_doc_limit = MAX_PER_DOC_API_CHARS
    if for_api and len(docs_list) > 1:
        per_doc_limit = max(1500, MAX_API_DOCUMENT_CHARS // len(docs_list))

    for doc in docs_list:
        name = doc.original_name
        text = (doc.extracted_text or '').strip()
        if not text:
            continue
        if for_api:
            text, was_cut = truncate_text(text, per_doc_limit, f'archivo «{name}»')
            truncated_any = truncated_any or was_cut
        blocks.append(f'### Archivo: {name}\n{text}')

    if not blocks:
        return '', False

    header = (
        'El profesor adjuntó los siguientes documentos. '
        'Usa su contenido para responder con precisión. '
        'Si algo no está en los archivos, indícalo.\n'
    )
    if for_api and truncated_any:
        header += (
            'Nota: el texto del documento fue recortado por tamaño; '
            'prioriza un resumen y menciona que el archivo completo es más largo.\n'
        )
    header += '\n'

    body = '\n\n---\n\n'.join(blocks)
    if for_api and len(body) > MAX_API_DOCUMENT_CHARS:
        body, was_cut = truncate_text(body, MAX_API_DOCUMENT_CHARS, 'documentos adjuntos')
        truncated_any = True

    return header + body, truncated_any


def file_icon_class(ext):
    icons = {
        '.pdf': 'bi-file-earmark-pdf',
        '.docx': 'bi-file-earmark-word',
        '.doc': 'bi-file-earmark-word',
        '.xlsx': 'bi-file-earmark-excel',
        '.xls': 'bi-file-earmark-excel',
        '.csv': 'bi-file-earmark-spreadsheet',
        '.txt': 'bi-file-earmark-text',
        '.md': 'bi-file-earmark-text',
    }
    return icons.get(ext, 'bi-file-earmark')
