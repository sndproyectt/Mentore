"""Generacion de archivos descargables para respuestas del chat IA."""
from __future__ import annotations

import csv
import io
import logging
import re
from dataclasses import dataclass

from django.core.files.base import ContentFile
from django.urls import reverse
from django.utils import timezone
from django.utils.text import get_valid_filename, slugify

from ai_assistant.models import GeneratedDocument

logger = logging.getLogger(__name__)

SUPPORTED_FORMATS = {'pdf', 'docx', 'xlsx', 'csv', 'txt', 'md'}

FORMAT_LABELS = {
    'pdf': 'PDF',
    'docx': 'Word',
    'xlsx': 'Excel',
    'csv': 'CSV',
    'txt': 'Texto',
    'md': 'Markdown',
}

FORMAT_ICONS = {
    'pdf': 'bi-file-earmark-pdf',
    'docx': 'bi-file-earmark-word',
    'xlsx': 'bi-file-earmark-excel',
    'csv': 'bi-file-earmark-spreadsheet',
    'txt': 'bi-file-earmark-text',
    'md': 'bi-markdown',
}

FORMAT_ALIASES = {
    'pdf': ('pdf',),
    'docx': ('word', 'docx', 'documento word'),
    'xlsx': ('excel', 'xlsx', 'hoja de calculo', 'hoja de cálculo'),
    'csv': ('csv',),
    'txt': ('txt', 'texto plano', 'archivo de texto'),
    'md': ('markdown', 'md'),
}

EXPORT_WORDS = (
    'genera', 'generame', 'créame', 'creame', 'crear', 'hazme', 'haz un',
    'exporta', 'descarga', 'descargable', 'pasame', 'pásame', 'convierte',
    'archivo', 'documento', 'informe',
)


@dataclass
class DocumentIntent:
    requested: bool
    file_format: str | None = None
    title: str | None = None


def detect_document_intent(message):
    text = (message or '').strip().lower()
    if not text:
        return DocumentIntent(False)

    file_format = None
    for fmt, aliases in FORMAT_ALIASES.items():
        if any(alias in text for alias in aliases):
            file_format = fmt
            break

    requested = bool(file_format and any(word in text for word in EXPORT_WORDS))
    return DocumentIntent(requested=requested, file_format=file_format, title=_guess_title(message, file_format))


def generate_document(user, file_format, content, title=None, source_prompt=''):
    file_format = (file_format or '').lower().strip()
    if file_format not in SUPPORTED_FORMATS:
        raise ValueError(f'Formato no soportado: {file_format}')

    title = title or 'Documento generado'
    filename = _filename(title, file_format)
    payload = _render_file(file_format, title, content or '')

    doc = GeneratedDocument(
        user=user,
        original_name=filename,
        file_format=file_format,
        source_prompt=(source_prompt or '')[:2000],
    )
    doc.file.save(filename, ContentFile(payload), save=False)
    doc.file_size = doc.file.size or len(payload)
    doc.save()
    logger.info("Documento IA generado: %s para %s", filename, user.username)
    return doc


def serialize_generated_document(doc, request=None):
    url = reverse('ai_assistant:generated_document_download', args=[doc.pk])
    if request is not None:
        url = request.build_absolute_uri(url)
    return {
        'id': doc.pk,
        'name': doc.original_name,
        'file_format': doc.file_format,
        'file_type': FORMAT_LABELS.get(doc.file_format, doc.file_format.upper()),
        'icon': FORMAT_ICONS.get(doc.file_format, 'bi-file-earmark'),
        'file_size': doc.file_size,
        'size_label': _format_size(doc.file_size),
        'created_at': timezone.localtime(doc.created_at).strftime('%d/%m/%Y %I:%M %p'),
        'download_url': url,
    }


def _guess_title(message, file_format):
    text = re.sub(r'\s+', ' ', (message or '')).strip()
    text = re.sub(r'\b(pdf|word|docx|excel|xlsx|csv|txt|markdown|md)\b', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\b(genera|generame|creame|crear|hazme|exporta|descarga|pasame|convierte|archivo|documento)\b', '', text, flags=re.IGNORECASE)
    words = [w for w in text.split() if len(w) > 2][:8]
    if words:
        return ' '.join(words).strip().capitalize()
    return f'Documento {FORMAT_LABELS.get(file_format, "IA")}'


def _filename(title, file_format):
    base = slugify(title or 'documento-generado') or 'documento-generado'
    timestamp = timezone.localtime().strftime('%Y%m%d_%H%M')
    return get_valid_filename(f'{base}_{timestamp}.{file_format}')


def _render_file(file_format, title, content):
    if file_format == 'docx':
        return _render_docx(title, content)
    if file_format == 'xlsx':
        return _render_xlsx(title, content)
    if file_format == 'csv':
        return _render_csv(content)
    if file_format == 'md':
        return f'# {title}\n\n{content.strip()}\n'.encode('utf-8')
    if file_format == 'txt':
        return f'{title}\n{"=" * len(title)}\n\n{content.strip()}\n'.encode('utf-8')
    if file_format == 'pdf':
        return _render_pdf(title, content)
    raise ValueError(f'Formato no soportado: {file_format}')


def _render_docx(title, content):
    import docx

    document = docx.Document()
    document.add_heading(title, level=1)
    for block in _paragraphs(content):
        if _looks_like_table_row(block):
            continue
        document.add_paragraph(block)

    table_rows = _extract_table_rows(content)
    if table_rows:
        document.add_heading('Tabla', level=2)
        table = document.add_table(rows=1, cols=len(table_rows[0]))
        table.style = 'Table Grid'
        for i, value in enumerate(table_rows[0]):
            table.rows[0].cells[i].text = value
        for row in table_rows[1:]:
            cells = table.add_row().cells
            for i, value in enumerate(row[:len(cells)]):
                cells[i].text = value

    stream = io.BytesIO()
    document.save(stream)
    return stream.getvalue()


def _render_xlsx(title, content):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = 'Resultado'
    rows = _extract_table_rows(content)
    if rows:
        for row in rows:
            ws.append(row)
    else:
        ws.append([title])
        ws.append([])
        for paragraph in _paragraphs(content):
            ws.append([paragraph])
    stream = io.BytesIO()
    wb.save(stream)
    wb.close()
    return stream.getvalue()


def _render_csv(content):
    output = io.StringIO()
    writer = csv.writer(output)
    rows = _extract_table_rows(content)
    if rows:
        writer.writerows(rows)
    else:
        writer.writerow(['contenido'])
        for paragraph in _paragraphs(content):
            writer.writerow([paragraph])
    return output.getvalue().encode('utf-8-sig')


def _render_pdf(title, content):
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.pdfgen import canvas

        stream = io.BytesIO()
        pdf = canvas.Canvas(stream, pagesize=letter)
        width, height = letter
        y = height - 72
        pdf.setFont('Helvetica-Bold', 14)
        pdf.drawString(72, y, _pdf_safe(title)[:80])
        y -= 28
        pdf.setFont('Helvetica', 10)
        for paragraph in _paragraphs(content):
            for line in _wrap(paragraph, 95):
                if y < 72:
                    pdf.showPage()
                    pdf.setFont('Helvetica', 10)
                    y = height - 72
                pdf.drawString(72, y, _pdf_safe(line))
                y -= 14
            y -= 6
        pdf.save()
        return stream.getvalue()
    except ImportError:
        return _render_minimal_pdf(title, content)


def _render_minimal_pdf(title, content):
    lines = [_pdf_safe(title), ''] + [_pdf_safe(line) for p in _paragraphs(content) for line in _wrap(p, 82)]
    objects = []
    stream_lines = ['BT', '/F1 10 Tf', '72 760 Td']
    for idx, line in enumerate(lines[:45]):
        if idx:
            stream_lines.append('0 -14 Td')
        stream_lines.append(f'({_escape_pdf_text(line)}) Tj')
    stream_lines.append('ET')
    content_stream = '\n'.join(stream_lines).encode('latin-1', errors='replace')

    objects.append(b'<< /Type /Catalog /Pages 2 0 R >>')
    objects.append(b'<< /Type /Pages /Kids [3 0 R] /Count 1 >>')
    objects.append(b'<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>')
    objects.append(b'<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>')
    objects.append(b'<< /Length ' + str(len(content_stream)).encode() + b' >>\nstream\n' + content_stream + b'\nendstream')

    out = io.BytesIO()
    out.write(b'%PDF-1.4\n')
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(out.tell())
        out.write(f'{i} 0 obj\n'.encode())
        out.write(obj)
        out.write(b'\nendobj\n')
    xref = out.tell()
    out.write(f'xref\n0 {len(objects) + 1}\n'.encode())
    out.write(b'0000000000 65535 f \n')
    for offset in offsets[1:]:
        out.write(f'{offset:010d} 00000 n \n'.encode())
    out.write(f'trailer << /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref}\n%%EOF'.encode())
    return out.getvalue()


def _paragraphs(content):
    text = (content or '').replace('\r\n', '\n').replace('\r', '\n')
    blocks = [block.strip() for block in re.split(r'\n{2,}', text) if block.strip()]
    if blocks:
        return blocks
    return [line.strip() for line in text.split('\n') if line.strip()]


def _extract_table_rows(content):
    rows = []
    for line in (content or '').splitlines():
        stripped = line.strip()
        if not _looks_like_table_row(stripped):
            continue
        parts = [part.strip() for part in stripped.strip('|').split('|')]
        if all(re.fullmatch(r':?-{2,}:?', part or '') for part in parts):
            continue
        rows.append(parts)
    return rows


def _looks_like_table_row(line):
    return bool(line and line.count('|') >= 2)


def _wrap(text, width):
    words = text.split()
    lines = []
    current = ''
    for word in words:
        if len(current) + len(word) + 1 > width:
            if current:
                lines.append(current)
            current = word
        else:
            current = f'{current} {word}'.strip()
    if current:
        lines.append(current)
    return lines or ['']


def _pdf_safe(text):
    return (text or '').replace('•', '-').replace('—', '-').replace('–', '-')


def _escape_pdf_text(text):
    return _pdf_safe(text).replace('\\', '\\\\').replace('(', '\\(').replace(')', '\\)')


def _format_size(num_bytes):
    if num_bytes < 1024:
        return f'{num_bytes} B'
    if num_bytes < 1024 * 1024:
        return f'{num_bytes // 1024} KB'
    return f'{num_bytes / (1024 * 1024):.1f} MB'
